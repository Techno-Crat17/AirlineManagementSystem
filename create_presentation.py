import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_presentation_docx():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Setup
    styles = doc.styles
    
    # Custom colors
    NAVY = RGBColor(28, 54, 115)
    CHARCOAL = RGBColor(51, 51, 51)
    CRIMSON = RGBColor(220, 53, 69)
    GRAY = RGBColor(108, 117, 125)
    
    # Title formatting helper
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(20)

    # Subtitle formatting helper
    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = GRAY
        p.paragraph_format.space_after = Pt(30)

    # Slide Heading formatting helper
    def add_slide_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.keep_with_next = True

    # Bullet point helper
    def add_bullet(text, bold_prefix="", indent=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
        p.paragraph_format.space_after = Pt(4)
        
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.font.name = 'Segoe UI'
            run_bold.font.size = Pt(11)
            run_bold.font.bold = True
            run_bold.font.color.rgb = CHARCOAL
            
        run_text = p.add_run(text)
        run_text.font.name = 'Segoe UI'
        run_text.font.size = Pt(11)
        run_text.font.color.rgb = CHARCOAL

    # ------------------ TITLE PAGE ------------------
    add_title("AIRLINE MANAGEMENT SYSTEM")
    add_subtitle("Presentation Preparation Guide & Speaker Notes\nTechnology Stack: Java Swing, JDBC, MySQL")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("This document is designed to help you prepare a comprehensive slideshow presentation about the Airline Management System project. It outlines the presentation slide-by-slide, giving you the Slide Content and corresponding Speaker Notes.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.color.rgb = CHARCOAL
    
    doc.add_page_break()

    # ------------------ SLIDE 1 ------------------
    add_slide_heading("Slide 1: Project Overview & Objectives")
    add_bullet(" A desktop-based ERP application designed to automate airline flight operations, passenger boarding, and database management.", "Project Definition: ")
    add_bullet(" Replace legacy manual ticket logging and Excel records with a relational database system.", "Primary Goal: ")
    add_bullet(" Admin Login, Passenger Profiles, Route Search, Booking, Cancellations, and Boarding Pass generation.", "Key Modules: ")
    add_bullet(" Clean UI design using standard font styling (Segoe UI), custom matching color schemes, and system Look-and-Feel styling.", "User Experience: ")
    
    p = doc.add_paragraph()
    run = p.add_run("Speaker Notes:\nStart by introducing the project as a desktop management system. Explain that the main target is to provide a single, unified database dashboard for airline ticketing agents. Mention that it aims to streamline onboarding, searching, reservation, and cancellation workflows under a secure, transaction-safe database environment.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ------------------ SLIDE 2 ------------------
    add_slide_heading("Slide 2: Technology Stack & Tools")
    add_bullet(" Java SE Development Kit (JDK 8+)", "Language & Runtime: ")
    add_bullet(" Java Swing (JFrame, JPanel, JTable, JTextField, JRadioButton, JComboBox, JMenuBar)", "GUI Library: ")
    add_bullet(" Java Database Connectivity (JDBC) API using MySQL Connector/J driver", "Connectivity: ")
    add_bullet(" MySQL (Relational Database Management System) for storing operational data", "Backend DB: ")
    add_bullet(" Apache Ant for building, compiling, and running the application", "Build Tool: ")
    add_bullet(" com.toedter.calendar (JDateChooser) for calendar date selections; net.proteanit.sql (DbUtils) for converting query ResultSets into visual Tables.", "Helper Libraries: ")
    
    p = doc.add_paragraph()
    run = p.add_run("Speaker Notes:\nDiscuss the choice of tools. Java Swing was selected for its robust framework and desktop-native rendering speed. JDBC provides secure connection pooling and Statement execution. MySQL handles large-volume data handling, indexing, and referential integrity constraints. Third-party helpers like DbUtils remove tedious boilerplate code for UI rendering.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)

    doc.add_page_break()

    # ------------------ SLIDE 3 ------------------
    add_slide_heading("Slide 3: System Architecture & Data Flow")
    add_bullet(" User interacts with custom Swing views (AddCustomer, BookFlight, etc.)", "1. Presentation Layer (GUI): ")
    add_bullet(" Connects GUI action events to SQL execution statements via Conn.java", "2. Data Access Layer (JDBC): ")
    add_bullet(" Runs structured statements to INSERT/SELECT/DELETE from tables", "3. Processing Logic: ")
    add_bullet(" Maintains relational tables, primary keys, and transaction states", "4. Database Layer (MySQL): ")
    
    p = doc.add_paragraph()
    run = p.add_run("Speaker Notes:\nExplain the flow of information. For example, when booking a flight, the user selects a source and destination. Swing triggers an ActionListener in BookFlight.java. It calls Conn.java, queries the database for matches, and returns flight information to the screen. Once confirmed, the data is pushed back into the reservation table, ensuring full separation of GUI rendering and backend persistence.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ------------------ SLIDE 4 ------------------
    add_slide_heading("Slide 4: Database Design & Schema")
    add_bullet(" stores admin users (username, password).", "login Table: ")
    add_bullet(" details of passengers (name, nationality, phone, address, aadhar [Primary Key], gender).", "passenger Table: ")
    add_bullet(" available schedules (f_code [Primary Key], f_name, source, destination).", "flight Table: ")
    add_bullet(" active reservations (PNR [Primary Key], ticket, aadhar, name, nationality, flightname, flightcode, src, des, ddate).", "reservation Table: ")
    add_bullet(" tracks canceled tickets (pnr, name, cancelno, fcode, date).", "cancel Table: ")
    
    p = doc.add_paragraph()
    run = p.add_run("Speaker Notes:\nDetail the schemas. Emphasize that 'aadhar' is used as a unique identifier for passenger profiles, preventing duplicates. Explain that the booking screen dynamically fetches user profiles based on this number. Point out that the 'cancel' table logs cancellation history which helps generate metrics on flight reliability and popularity.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)

    doc.add_page_break()

    # ------------------ SLIDE 5 ------------------
    add_slide_heading("Slide 5: Key Feature Implementation Highlights")
    add_bullet(" Secure, password-hidden form querying the login database.", "Admin Login Authentication: ")
    add_bullet(" DocumentFilter (DigitLimitFilter) locks fields to numerical inputs and limits string length (e.g., exactly 12 digits for Aadhar, 10 for Phone) preventing user error.", "Input Sanitation & Restrictions: ")
    add_bullet(" Selecting source/destination cities updates the Flight Name and Code automatically, avoiding manual code entry mistakes.", "Dynamic Flight Lookup: ")
    add_bullet(" Auto-fetches name, gender, nationality, and address details dynamically as soon as a registered Aadhar number is entered.", "Passenger Profile Fetching: ")
    add_bullet(" Automatically prints all ticket parameters onto a clean UI layout for flight agents and customers.", "On-demand Boarding Pass: ")
    add_bullet(" Moves reservation entry to the cancellation log and clears the active booking queue in a single transaction step.", "One-click Cancellation: ")
    
    p = doc.add_paragraph()
    run = p.add_run("Speaker Notes:\nWalk the audience through the technical highlights. Explain that usability is key. For example, DigitLimitFilter prevents users from typing letters in phone number fields. Automatic profile fetching cuts down data-entry time during booking. The instant boarding pass generation right after booking improves agent workflow.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ------------------ SLIDE 6 ------------------
    add_slide_heading("Slide 6: Complex Database Queries (Analytics)")
    add_bullet(" Finds busiest flight connections based on reservation counts.", "Reservation Volumes: ")
    add_bullet(" Tracks cancellation ratios per flight to identify low-performing schedules.", "Flight Cancellation Rates: ")
    add_bullet(" Analyzes travel hubs (departure and arrival cities) to rank top locations.", "Route Optimization: ")
    add_bullet(" Identifies family or duplicate profiles sharing the same phone contact details using GROUP_CONCAT.", "Contact Audit: ")
    add_bullet(" Leverages UNION ALL to compile unified reservation and cancellation statements in a single search query.", "Passenger Ledgers: ")
    
    p = doc.add_paragraph()
    run = p.add_run("Speaker Notes:\nHighlight how backend data is utilized for decision support. Using complex queries (like those in complex_queries.sql), management can see cancellation rates, gender distribution, and passenger booking counts. Mention that the database supports analytics like finding top routes to optimize scheduling and revenue.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)

    doc.add_page_break()

    # ------------------ SLIDE 7 ------------------
    add_slide_heading("Slide 7: Conclusion & Future Scope")
    add_bullet(" Successfully creates a fully responsive, desktop-based client database system.", "Summary: ")
    add_bullet(" Streamlines passenger data management and booking operations.", "Key Benefit: ")
    add_bullet(" Standardizing JDBC operations prevents connection leakage; utilizing modern Segoe UI styling ensures a professional feel.", "Development Lesson: ")
    add_bullet(" Migrate to a multi-tiered Web/REST API framework, implement seat selection grids, and add email/SMS booking notifications.", "Future Upgrades: ")
    
    p = doc.add_paragraph()
    run = p.add_run("Speaker Notes:\nWrap up the presentation by highlighting project success. Conclude with lessons learned about database constraints and UI validation. Invite questions from the audience, pointing out that this architecture serves as a stable base for future cloud integration or web-based ticketing portals.")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)

    doc.save("Airline_Management_System_Presentation_Prep.docx")
    print("Created Airline_Management_System_Presentation_Prep.docx successfully.")

if __name__ == "__main__":
    create_presentation_docx()
